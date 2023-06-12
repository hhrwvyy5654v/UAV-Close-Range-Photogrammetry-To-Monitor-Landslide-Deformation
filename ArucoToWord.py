'''
Description: 获取Aruco码中心点的像素坐标
FilePath: \ArucoToWord.py
Author: hhrwvyy5654v huang_rongquan@outlook.com
Date: 2023-05-15 20:48:08
LastEditors: hhrwvyy5654v huang_rongquan@outlook.com
LastEditTime: 2023-06-12 09:53:22
Copyright (c) 2023 by hhrwvyy5654v , All Rights Reserved. 
'''

# 导包
import cv2
import numpy as np
from docx import Document
from docx.shared import Inches


# 创建Aruco码,输入参数为:行数、列数、Aruco码字典、标记长度、标距间距、码标
def create_aruco_markers(rows, cols, aruco_dict, marker_length, marker_separation, marker_ids):
    aruco_params = cv2.aruco.DetectorParameters_create()

    # 创建3D对象点数组：包含4个点的数组，用于定义标记的形状和大小
    obj_points = np.zeros((4, 3), dtype=np.float32)
    obj_points[:, :2] = np.array([[0, 0], [marker_length, 0], [
                                 marker_length, marker_length], [0, marker_length]])
    obj_points_array = [obj_points] * (rows * cols)

    # 创建Aruco板:使用上一步中定义的对象点数组、字典和记ID创建Aruco板
    aruco_board = cv2.aruco.Board_create(objPoints=obj_points_array,
                                         dictionary=aruco_dict,
                                         ids=marker_ids)
    img_size = (cols * (marker_length + marker_separation) + marker_separation,
                rows * (marker_length + marker_separation) + marker_separation)
    
    # 创建空白图像:根据行数、列数、标记长度和标记间距创建一个空白图像
    aruco_image = np.full(img_size, 255, dtype=np.uint8)
    
    # 绘制标记:遍历Aruco板中的每个标记,使用cv2.aruco.drawMarker函数绘制标记,并将其放置在空白图像中的相应位置
    for i, obj_points in enumerate(aruco_board.objPoints):
        marker_image = cv2.aruco.drawMarker(
            aruco_dict, marker_ids[i][0], marker_length)
        x = marker_separation + (i % cols) * \
            (marker_length + marker_separation)
        y = marker_separation + (i // cols) * \
            (marker_length + marker_separation)
        aruco_image[y:y + marker_length, x:x + marker_length] = marker_image
    
    # 返回包含所有标记的Aruco图像
    return aruco_image


# 保存Aruco码标记到Word文件
def save_aruco_markers_to_word(doc, aruco_image, image_name,doc_name):
    # 使用cv2.imwrite函数将Aruco标记图像保存到指定的图像文件中
    cv2.imwrite(image_name, aruco_image)
    # 使用doc.add_picture函数将保存的图像文件插入到Word文档中，并指定图像的宽度和高度
    doc.add_picture(image_name, width=Inches(8.27), height=Inches(11.69))
    # 使用doc.save函数将修改后的Word文档保存到指定的文件中
    doc.save(doc_name)


# 设置Aruco标记的参数
rows = 2    # 行数
cols = 3    # 列数
aruco_dict = cv2.aruco.Dictionary_get(cv2.aruco.DICT_6X6_250) # 字典
marker_length = 50  # 码标长度(单位：毫米)
marker_separation = 50  # 标记间距(单位：毫米)
marker_ids = np.array([[0], [1], [2], [3], [4], [5]]).reshape(rows, cols) # 标记ID

# 将标记长度和标记间距从毫米转换为像素,假设每英寸96像素
dpi = 96  
marker_length_pixels = int(marker_length * dpi / 25.4)
marker_se_pixels = int(marker_separation * dpi / 25.4)

# 生成Aruco码标记图像
aruco_image = create_aruco_markers(
    rows, cols, aruco_dict, marker_length_pixels, marker_separation_pixels, marker_ids)

# 保存Aruco码到Word文档中
doc = Document()
save_aruco_markers_to_word(
    doc, aruco_image, './ArucoGenerated/aruco_markers.png', './ArucoGenerated/aruco_markers.docx')
